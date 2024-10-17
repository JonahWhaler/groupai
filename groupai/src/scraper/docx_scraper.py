from docx import Document
import zipfile
import os
import re
from io import StringIO, BytesIO
from typing import Optional, List
from .base_scraper import BaseScraper
from .image_scraper import ImageToMarkdownScraper


class DocxToMarkdownScraper(BaseScraper):
    def __init__(self, **kwargs):
        self.image_interpretor: Optional[ImageToMarkdownScraper] = kwargs.get(
            "image_interpretor", None)
        self.tmp_directory = kwargs.get("tmp_directory", None)
        os.makedirs(self.tmp_directory, exist_ok=True)

    def get_cell_content_with_formatting(self, cell):
        content = StringIO()
        for para in cell.paragraphs:
            for run in para.runs:
                if run.bold:
                    content.write(f"**{run.text}**")
                elif run.italic:
                    content.write(f"*{run.text}*")
                else:
                    content.write(run.text)
            content.write('\n')
        return content.getvalue().strip()

    def temporary_store_images(self, input_path: str) -> List[str]:
        images = []
        with zipfile.ZipFile(input_path, 'r') as docx:
            # Iterate through the files in the archive
            for file in docx.namelist():
                # Check if the file is an image located in 'word/media/'
                if file.startswith('word/media/'):
                    # Extract the image
                    image_data = docx.read(file)
                    image_filename = os.path.join(
                        self.tmp_directory, os.path.basename(file))

                    # Write the image to the output directory
                    with open(image_filename, 'wb') as img_file:
                        img_file.write(image_data)
                    images.append(image_filename)
        return images

    def to_markdown(self, input_path: str):
        doc = Document(input_path)
        markdown = []

        # Iterate through all elements in the document
        for para in doc.paragraphs:
            # Handle headings
            if para.style.name.startswith('Heading'):
                level = int(re.search(r'\d+', para.style.name).group(0))
                markdown.append(f"{'#' * level} {para.text}")
            else:
                # Handle regular paragraphs and text
                if para.text.strip():
                    markdown.append(para.text)
            markdown.append('\n')
        
        # Extract tables
        for table in doc.tables:
            markdown.append('\n')
            for row_index, row in enumerate(table.rows):
                row_data = [
                    f"| {self.get_cell_content_with_formatting(cell)} " for cell in row.cells]
                markdown.append(''.join(row_data) + '|')
                # Create a separator for header row (assuming first row is header)
                if row_index == 0:
                    header_separator = ''.join(['| --- ' for _ in row.cells])
                    markdown.append(header_separator + '|')
            # Add a newline after the table
            markdown.append('\n')

        # Extract images
        markdown.append(f"\n# Images\n")
        # for rel in doc.part.rels.values():
        #     if "image" in rel.target_ref:
        #         markdown.append(f"![]({rel.target_ref})")
        # markdown.append('\n')
        images = self.temporary_store_images(input_path)
        for image in images:
            image_description = self.image_interpretor.to_markdown(
                image_path=image, caption="This is an attachment found in a pdf file.")
            markdown.append(f"## {image}\n{image_description}\n")
            os.remove(image)
        # Return markdown as a single joined string
        return '\n'.join(markdown)

    def __call__(
        self, input_path: Optional[str] = None, buffer: Optional[BytesIO] = None,
        identifier: Optional[str] = None, caption: Optional[str] = None,
        output_path: Optional[str] = None
    ) -> str:
        assert input_path is not None and input_path != "" and input_path[-4:] == "docx"

        output_markdown = self.to_markdown(input_path)
        if output_path is not None:
            self.save_markdown(output_markdown, output_path)
        return output_markdown


if __name__ == '__main__':
    # Usage
    OPENAI_API_KEY = ""
    model = "gpt-4o-mini-2024-07-18"
    image_scraper = ImageToMarkdownScraper(OPENAI_API_KEY=OPENAI_API_KEY, vision_model=model)
    
    converter = DocxToMarkdownScraper(image_interpretor=image_scraper, tmp_directory='./asset')
    converter(input_path='example.docx', output_path='./asset/docx.md')
